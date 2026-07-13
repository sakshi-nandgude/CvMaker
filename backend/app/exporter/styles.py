from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


def configure_document(document):
    """
    Configure page layout and default styles.
    """

    section = document.sections[0]

    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = document.styles["Normal"]

    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)


def add_heading(document, text):
    """
    Section Heading
    """

    paragraph = document.add_paragraph()

    run = paragraph.add_run(text.upper())

    run.bold = True
    run.font.size = Pt(12)

    paragraph.space_before = Pt(8)
    paragraph.space_after = Pt(4)

    return paragraph


def add_name(document, name):
    """
    Candidate Name
    """

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run(name)

    run.bold = True
    run.font.size = Pt(18)

    paragraph.space_after = Pt(2)


def add_title(document, title):
    paragraph = document.add_paragraph()

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run(title)

    run.font.size = Pt(11)

    paragraph.space_after = Pt(4)


def add_contact(document, text):
    paragraph = document.add_paragraph()

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run(text)

    run.font.size = Pt(9)

    paragraph.space_after = Pt(8)